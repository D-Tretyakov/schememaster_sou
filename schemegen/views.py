from django.http import HttpResponse, Http404
from django.shortcuts import render
from django.conf import settings

from rest_framework import viewsets
from rest_framework.response import Response
from rest_framework.decorators import api_view

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate
from datetime import datetime
import json
import re
import os

from .models import Template, TextAlias

def index(request):
    if not request.session or not request.session.session_key:
        request.session.save()
    context = {'full_path':  request.build_absolute_uri()}
    return render(request, 'schemegen/index.html', context)

# def get_tree(request, tree_id):
#     tree_ids = {tree.id: tree.name for tree in Tree.objects.all()}
#     tree = Tree.objects.get(id=tree_id)
#     context = {'tree': tree, 'tree_ids': tree_ids}
#     return render(request, 'schemegen/tree.html', context)

# def convert(request, tree_id):
#     print(request.POST)
#     result = [Variant.objects.get(id=int(request.POST[f'variant_choice{c.id}'])).text_repr
#             for c in Tree.objects.get(id=tree_id).choice_set.all()]
#     text = Tree.objects.get(id=tree_id).schema_set.all()[0].text_repr.format(*result)

#     return HttpResponse(text.replace('\n', '<br>'))

def convertion(ans, session_key):
    document = Document()
    fname = f'./claims/demo_{session_key}.docx'
    document.save(fname)  
    p = document.add_paragraph()
    p.add_run('В___________________').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = document.add_paragraph()
    p.add_run("Адрес:_____________________").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT    
    for i in ans[0]:
        p = document.add_paragraph()
        p.add_run(i).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for j in ans[1]:
        res = re.split('\n',j)
        p = document.add_paragraph()
        p.add_run('Истец:').bold = True
        for i in res:
            p.add_run(i)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p = document.add_paragraph()
    if ans[2][0] != '':
        k = 0
        for i in ans[2][1:]:
            res = re.split(r'[1]\d*.|\n[2,3]\d*.',i)
            res.pop(0)
            part1 = re.split('\n',res[0])
            p.add_run('Представитель Истца:').bold = True
            ans[2][k]=res[1]
            for j in part1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p = document.add_paragraph() 
            k+=1
    for i in ans[3]:
        res = re.split('\n',i)
        p = document.add_paragraph()
        p.add_run('Ответчик:').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for j in res:            
            p.add_run(j)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT      
            p = document.add_paragraph()
    if ans[4][0] != '':
        for i in ans[4][1:]:
            p.add_run('Третье лицо: ').bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            res = re.split('\n',i)
            for j in res:
                p.add_run(j)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
                p = document.add_paragraph()
    if ans[5] != ['']:
        p = document.add_paragraph()
        p.add_run('{{price_isk}}').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if ans[6] != ['']:
        p = document.add_paragraph()
        p.add_run('Государственная пошлина: ').bold = True
        p.add_run('{{poshlina}}')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = document.add_paragraph()
    p.add_run("Исковое заявление").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    k = 0
    l = []
    for i in ans[7]:
        res = re.split(r'[1]\d*.|\n[2,3]\d*.',i)
        res.pop(0)
        l.append(res)
        if i == ans[7][-1]:
            ans[7] = l
            break
    l=[]
    for i in ans[8]:
        res = re.split(r'[1]\d*.|\n[2,3]\d*.',i)
        res.pop(0)
        l.append(res)
        if i == ans[8][-1]:
            ans[8] = l
            break
    p = document.add_paragraph()
    s = ''
    for i in ans[7]:
        s += i[0]+', '
    s = s[:-2]
    p.add_run('о ' + s).italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph("ПРОШУ:")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    k = 0
    for i in ans[7]:
        k+=1
        document.add_paragraph('%i.' %k + i[1])
    for i in ans[8]:
        k+=1
        document.add_paragraph('%i.' %k + i[0])
    if ans[11]!= ['']:
        k+=1
        document.add_paragraph('%i.' %k + '{{potrebiteli}}')
    p = document.add_paragraph()
    p.add_run('Приложение:').bold = True
    if ans[6] != ['']:
        document.add_paragraph('Платежное поручение №___ от «__»______ ____ г., подтверждающее уплату государственной пошлины.\n\n« » ________ _____г. _____________ (________________)', style = 'List Number')
    document.add_paragraph('Копия уведомления о вручении или иные документы, подтверждающие направление другим лицам, участвующим в деле, копий искового заявления и приложенных к нему документов, которые у других лиц, участвующих в деле, отсутствуют;', style = 'List Number')
    document.add_paragraph('Иные документы, на которых Истец обосновывает свои требования;', style = 'List Number')
    if ans[2][0]!= '':
        document.add_paragraph('{{complainant}}', style = 'List Number')
    for i in ans[7]:
        res = re.split('\n',i[2])
        document.add_paragraph(res[0], style = 'List Number')
        if len(res) > 1:
            for j in res[1:]:
                document.add_paragraph(j)
    l = []
    for i in ans[8]:
        l.append(i[1])
    res = set(l)
    for i in res:
        document.add_paragraph(res, style = 'List Number')
    if ans[9] != ['']:
        document.add_paragraph('{{regulation}}', style = 'List Number')
    if ans[10] != ['']:
        document.add_paragraph('{{peace}}', style = 'List Number')
    document.add_paragraph('« » ________ _____г. \t\t\t\t\t\t_____________ (________________)')
    document.save(fname)
    document = DocxTemplate(fname)
    context = { 'complainant':ans[2][0], 'price_isk':ans[5][0],'poshlina':ans[6][0], 'regulation':ans[9][0], 'peace':ans[10][0], 'potrebiteli':ans[11][0]}
    document.render(context)
    document.save(fname)

def get_text(request):
    template = Template.objects.first()
    header = "<div style='text-align: right;'>" + template.header + "</div>"
    body = "<div style='text-align: center;'>" + template.body + "</div>"
    footer = template.footer

    schema = header + body + footer
    
    req = dict(request.POST)
    # print(req)
    # if req['choice-7'][0] in ['c7-v6', 'c7-v7', 'c7-v8']:
    #     req['choice-7'].extend(req['choice-7-1'])
    # del req['choice-7-1']



    res = {}
    for choice in req:
        if len(choice.split('-')) != 2:
            continue

        i = int(choice.split('-')[1])
        answer = []

        # if choice == 'choice-7' and len(req[choice]) == 2:
        #     ans0 = TextAlias.objects.get(html_id=req[choice][0]).text
        #     ans1 = TextAlias.objects.get(html_id=req[choice][1]).text
        #     answer.append(ans0+'\n'+ans1)
            # res[i] = answer
        if choice == 'choice-7' and 'c7-v2' in req[choice]:
            op1 = req.get('choice-7-op1')
            op2 = req.get('choice-7-op2')
            op3 = req.get('choice-7-op3')
            op4 = req.get('choice-7-op4')

            if op1 is not None:
                answer.append(TextAlias.objects.get(html_id=op1[0]).text)
            if op2 is not None and op3 is not None:
                ans0 = TextAlias.objects.get(html_id=op2[0]).text
                ans1 = TextAlias.objects.get(html_id=op3[0]).text
                answer.append(ans0 + '\n' + ans1)
            if op4 is not None:
                answer.append(TextAlias.objects.get(html_id=op4[0]).text)

            req['choice-7'].remove('c7-v2')
        
        for j in req[choice]:
            answer.append(TextAlias.objects.get(html_id=j).text)
        
        res[i] = answer
    
    print(res)
    # ans = [res[key] for key in sorted(res.keys())]    
    # text = schema.format(*ans)
    convertion(res, request.session.session_key)
    # return HttpResponse(text.replace('\n', '<br>'))
    return HttpResponse("Success")

def download(request):
    session_key = request.session.session_key
    os.system(f'doc2pdf --output=./claims/template_{session_key} ./claims/demo_{session_key}.docx')

    path = f'./claims/template_{session_key}.pdf'
    file_path = os.path.join(settings.MEDIA_ROOT, path)
    print(file_path)
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/pdf")
            response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
            return response
    raise Http404

def download_doc(request):
    # os.system('doc2pdf --output=template demo.docx')
    session_key = request.session.session_key
    path = f'./claims/demo_{session_key}.docx'
    file_path = os.path.join(settings.MEDIA_ROOT, path)
    dt = datetime.now()
    # print(file_path)
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/msdocx")
            response['Content-Disposition'] = 'inline; filename=' + f'Generated_claim_{dt.strftime("%Y%m%d%H%M%S")}.docx'
            return response
    raise Http404

# @api_view(['POST'])
# def pretty_print(request):   
#     req = dict(request.POST)
#     del req['csrfmiddlewaretoken']
    
#     res = {}
#     for choice in req:
#         i = int(choice.split('-')[1])
#         answer = []
#         for j in req[choice]:
#             answer.append(TextAlias.objects.get(html_id=j).text)
#             res[i] = answer 
#     ans = []
#     for key in sorted(res.keys()):
#       ans.append(res[key])

#     return Response({
#         'req': req,
#         'res':res,
#         'ans': ans,
#     })

# def get_text(request):
#     template = Template.objects.first()
#     header = "<div style='text-align: right;'>" + template.header + "</div>"
#     body = "<div style='text-align: center;'>" + template.body + "</div>"
#     footer = template.footer

#     schema = header + body + footer
    
#     req = dict(request.POST)
#     res = {}
#     for choice in req:
#         i = int(choice.split('-')[1])
#         res[i] = TextAlias.objects.get(html_id=req[choice][0]).text
   
#     ans = [res[key] for key in sorted(res.keys())]
#     ans += ['1. Копия доверенности или иного документа, подтверждающего полномочия представителя' if res[2] else '']
#     ans += ['2. Платежное поручение №___ от «__»______ ____ г., подтверждающее уплату государственной пошлины.\n\n« » ________ _____г. _____________ (________________)'
#             if res[6] != 'не взимается' else '']
#     text = schema.format(*ans)

#     return HttpResponse(text.replace('\n', '<br>'))

# @api_view()
# def pass_func(request):
#     template = Template.objects.first()
#     header = "<div style='text-align: right;'>" + template.header + "</div>"
#     body = "<div style='text-align: center;'>" + template.body + "</div>"
#     footer = template.footer

#     schema = header + body + footer

#     return Response({'text': schema})